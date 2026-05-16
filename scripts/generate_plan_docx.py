#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成液冷市场突破计划 Word 文档
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """设置表格单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = parse_xml(r'<w:tcBorders %s/>' % nsdecls('w'))
        tcPr.append(tcBorders)

def add_heading_custom(doc, text, level=1):
    """添加自定义标题"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_table_with_style(doc, data, has_header=True):
    """添加带样式的表格"""
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Light Grid Accent 1'
    
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_text)
            # 设置字体
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
            # 表头加粗
            if has_header and i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
    
    return table

# 创建文档
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 标题
title = doc.add_heading('液冷软管及快接头市场突破计划', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 文档信息
info_table_data = [
    ['文件版本：', 'V1.0', '制定日期：', '2026年5月'],
    ['制定人：', '总经理', '适用范围：', '内部传阅']
]
info_table = doc.add_table(rows=2, cols=4)
for i, row_data in enumerate(info_table_data):
    for j, text in enumerate(row_data):
        cell = info_table.rows[i].cells[j]
        cell.text = text
        if j % 2 == 0:  # 标签列加粗
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

doc.add_paragraph()

# 一、战略背景与目标
add_heading_custom(doc, '一、战略背景与目标', 1)

add_heading_custom(doc, '1.1 市场判断', 2)
doc.add_paragraph(
    '根据前期市场调研，AI数据中心（AIDC）建设浪潮带动液冷系统需求快速增长，液冷软管及快接头作为液冷系统的关键连接部件，'
    '正处于市场导入期向成长期过渡的窗口。当前市场竞争格局尚未固化，具备专业方案能力的供应商稀缺，是我方切入的战略时机。'
)

add_heading_custom(doc, '1.2 核心战略定位', 2)
p = doc.add_paragraph()
p.add_run('不以单品销售为目标，以"液冷系统连接解决方案提供商"为市场定位。').bold = True
doc.add_paragraph(
    '快接头与软管是产品入口，行业解决方案能力是竞争壁垒，稳定的供应链是交付保障。三者缺一不可。'
)

add_heading_custom(doc, '1.3 五月总体目标', 2)
goal_data = [
    ['目标维度', '具体指标'],
    ['团队', '核心售前团队框架搭建完成，技术支撑岗位到位或完成招募'],
    ['市场', '线上营销体系（网站专题页+公众号）上线运营'],
    ['供应链', '完成UDQ供应商深度对接，形成供应商能力评估档案'],
    ['业务', '完成不少于10家目标客户的初步接触，建立商机管道']
]
add_table_with_style(doc, goal_data)

doc.add_page_break()

# 二、团队能力架构
add_heading_custom(doc, '二、团队能力架构', 1)

add_heading_custom(doc, '2.1 团队能力架构图', 2)
doc.add_paragraph(
    '┌─────────────────────────────────────────────────────────┐\n'
    '│                    总经理（统筹决策）                      │\n'
    '│              战略客户关系 · 资源协调 · 方案拍板             │\n'
    '└──────────────────────┬──────────────────────────────────┘\n'
    '                       │\n'
    '        ┌──────────────┼──────────────┐\n'
    '        │              │              │\n'
    '┌───────▼──────┐ ┌─────▼──────┐ ┌────▼────────────┐\n'
    '│  技术支撑模块  │ │ 商务拓展模块 │ │  市场营销模块    │\n'
    '│              │ │            │ │                 │\n'
    '│ 液冷技术工程师 │ │ 售前项目经理 │ │  内容运营专员    │\n'
    '│  （1-2人）   │ │  （1人）    │ │   （1人）        │\n'
    '│              │ │            │ │                 │\n'
    '│ · 方案设计    │ │ · 商机跟进  │ │ · 网站/公众号    │\n'
    '│ · 参数匹配    │ │ · 客户对接  │ │ · 内容策划生产   │\n'
    '│ · 技术答疑    │ │ · 项目推进  │ │ · 线上获客      │\n'
    '│ · 供应商评估  │ │ · 合同跟进  │ │ · 数据分析      │\n'
    '└──────────────┘ └────────────┘ └─────────────────┘\n'
    '        │\n'
    '        │  【集团总部技术支持（协同）】\n'
    '        │  · 液冷技术能力建设\n'
    '        │  · 供应商技术鉴别支持\n'
    '        │  · 行业标准与规范输入\n'
    '        └──────────────────────────',
    style='Normal'
).paragraph_format.line_spacing = 1.0

doc.add_paragraph(
    '说明：当前阶段团队规模精简，以"小团队、强协同"为原则。技术支撑模块与集团总部技术团队保持紧密协同，'
    '避免重复建设。各模块负责人直接向总经理汇报，确保决策链条短、响应速度快。'
)

doc.add_page_break()

# 三、技术支撑岗位能力要求
add_heading_custom(doc, '三、技术支撑岗位能力要求', 1)
doc.add_paragraph('本节内容可直接提供给HR或猎头作为招募依据。').italic = True

add_heading_custom(doc, '3.1 液冷技术工程师（优先级：最高）', 2)
doc.add_paragraph('岗位定位：').bold = True
doc.add_paragraph('团队核心技术支柱，负责对外的技术方案输出与对内的供应商技术评估。')

add_heading_custom(doc, '必备能力要求', 3)
doc.add_paragraph('专业知识（硬性要求）', style='Heading 4')

knowledge_data = [
    ['能力项', '具体要求'],
    ['液冷系统原理', '熟悉数据中心液冷系统架构，包括CDU（冷量分配单元）、冷板式、浸没式等主流方案'],
    ['流体力学基础', '理解压降、流量、流速等核心参数，能进行基础水力计算'],
    ['快接头与软管知识', '熟悉主流品牌（Colder、CEJN、Parker、UDQ等）产品系列，了解材质、压力等级、温度范围、接口标准'],
    ['行业标准', '了解ASHRAE、ETSI、GB/T等数据中心相关标准中对液冷系统的要求'],
    ['方案设计能力', '能根据客户机柜功率密度、散热需求，输出液冷连接方案选型建议书']
]
add_table_with_style(doc, knowledge_data)

doc.add_paragraph()
doc.add_paragraph('工作经验要求', style='Heading 4')
doc.add_paragraph('• 3年以上液冷相关工作经验，优先考虑以下背景：')
doc.add_paragraph('  - 数据中心液冷系统集成商（技术/售前岗）', style='List Bullet 2')
doc.add_paragraph('  - 服务器/机柜厂商液冷研发或应用工程岗', style='List Bullet 2')
doc.add_paragraph('  - 液冷部件（泵、阀、管路）供应商技术支持岗', style='List Bullet 2')
doc.add_paragraph('• 有AIDC项目实施或交付经验者优先')
doc.add_paragraph('• 有主流液冷快接头品牌产品使用或评测经验者优先')

doc.add_paragraph()
doc.add_paragraph('软性能力要求', style='Heading 4')
doc.add_paragraph('• 具备良好的客户沟通能力，能将技术内容转化为客户可理解的语言')
doc.add_paragraph('• 能独立撰写技术方案文档、选型报告')
doc.add_paragraph('• 具备一定的市场敏感度，能从技术角度识别客户痛点')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('薪资参考范围：').bold = True
p.add_run('20,000 - 35,000元/月（视经验和背景浮动）')

doc.add_paragraph()
doc.add_paragraph('招募渠道建议：', style='Heading 4')
doc.add_paragraph('1. 猎头定向搜索：华为数字能源、中兴、曙光、超聚变等厂商的液冷技术岗')
doc.add_paragraph('2. BOSS直聘/猎聘：关键词"液冷工程师"、"数据中心热管理"、"CDU"')
doc.add_paragraph('3. 行业展会（数据中心展、IDC峰会）现场接触')
doc.add_paragraph('4. 短期过渡方案：以技术顾问/兼职形式先行合作，降低招募时间风险')

doc.add_page_break()

# 四、供应商体系构建与鉴别能力建设
add_heading_custom(doc, '四、供应商体系构建与鉴别能力建设', 1)

add_heading_custom(doc, '4.1 背景说明', 2)
doc.add_paragraph(
    '当前团队在液冷技术领域尚处于能力建设初期，集团总部正在同步推进技术能力构建。在此阶段，建立一套结构化的供应商鉴别与评估体系，'
    '是确保对外方案可信度和交付稳定性的关键基础工作。'
)
doc.add_paragraph(
    '本节旨在提供一套可操作的供应商鉴别框架，使团队在技术能力尚未完全到位的情况下，仍能对供应商进行有效评估和筛选。'
)

add_heading_custom(doc, '4.2 供应商鉴别能力建设路径', 2)

add_heading_custom(doc, '第一步：建立评估知识基础（第1-2周）', 3)
doc.add_paragraph('在与供应商正式接触前，团队需掌握以下基础知识：')

knowledge_base_data = [
    ['知识点', '重要性', '学习来源'],
    ['液冷快接头的主要技术参数（工作压力、温度范围、流量系数Cv值、泄漏率）', '★★★', '主流品牌技术手册'],
    ['常见材质及适用场景（不锈钢316L、黄铜、PVDF、PEEK）', '★★★', '供应商产品手册'],
    ['主流国际品牌及其定位（Colder/Stäubli/CEJN/Parker）', '★★', '行业资料'],
    ['数据中心液冷对快接头的特殊要求（无泄漏、低压降、免工具操作）', '★★★', 'ASHRAE TC9.9指南'],
    ['国内外认证标准（CE、UL、RoHS、IP等级）', '★★', '认证机构官网']
]
add_table_with_style(doc, knowledge_base_data)

doc.add_page_break()

add_heading_custom(doc, '第二步：供应商评估维度框架', 3)

doc.add_paragraph('维度一：产品技术能力（权重40%）', style='Heading 4')
tech_eval_data = [
    ['评估项', '评估方法', '关注要点'],
    ['产品覆盖范围', '索取完整产品目录', '是否覆盖液冷专用系列；规格是否满足主流服务器机柜需求'],
    ['技术参数真实性', '要求提供第三方检测报告', '工作压力、温度范围、泄漏率数据是否有实验室背书'],
    ['产品认证情况', '查验认证证书原件', 'CE/UL认证是否在有效期内；认证范围是否覆盖目标产品'],
    ['研发能力', '了解研发团队规模及近3年新品发布情况', '是否具备定制开发能力；是否跟随液冷行业趋势迭代产品'],
    ['样品测试', '索取样品，委托集团技术团队或第三方机构测试', '实测参数与标称参数的偏差；装配手感与耐久性']
]
add_table_with_style(doc, tech_eval_data)

doc.add_paragraph()
doc.add_paragraph('维度二：供应稳定性（权重30%）', style='Heading 4')
stability_data = [
    ['评估项', '评估方法', '关注要点'],
    ['生产能力', '参观工厂或要求提供产能证明', '月产能是否满足预期订单规模；是否有产能扩充计划'],
    ['交货周期', '明确常规订单与紧急订单的交期', '常规交期是否在15个工作日内；紧急订单是否有加急通道'],
    ['库存策略', '了解其成品库存政策', '主力SKU是否有现货库存；安全库存水位如何'],
    ['原材料供应链', '了解关键原材料来源', '是否依赖单一原材料供应商；是否有原材料价格波动应对机制'],
    ['历史交付记录', '要求提供近1年交货准时率数据', '准时率是否高于95%']
]
add_table_with_style(doc, stability_data)

doc.add_paragraph()
doc.add_paragraph('维度三：商务合作条件（权重20%）', style='Heading 4')
business_data = [
    ['评估项', '评估方法', '关注要点'],
    ['授权与代理政策', '明确合作模式（经销/代理/直采）', '是否提供区域保护；价格体系是否透明稳定'],
    ['价格竞争力', '横向比较同类产品报价', '与国际品牌的价格差异；批量折扣政策'],
    ['账期与付款条件', '谈判付款方式', '是否支持账期；首单政策如何'],
    ['售后支持', '了解退换货政策及质保条款', '质保期限；不良品处理流程与时效']
]
add_table_with_style(doc, business_data)

doc.add_paragraph()
doc.add_paragraph('维度四：企业资质与信誉（权重10%）', style='Heading 4')
credential_data = [
    ['评估项', '评估方法', '关注要点'],
    ['企业规模与历史', '查询工商信息（天眼查/企查查）', '注册资本、成立年限、股权结构是否稳定'],
    ['客户背书', '要求提供参考客户名单', '是否有知名AIDC或服务器厂商客户；可联系参考客户核实'],
    ['行业口碑', '通过行业圈子侧面了解', '是否有质量纠纷或交货违约记录'],
    ['知识产权', '查询专利情况', '核心产品是否有自主专利；是否存在侵权风险']
]
add_table_with_style(doc, credential_data)

doc.add_page_break()

add_heading_custom(doc, '第三步：供应商评估执行流程', 3)
doc.add_paragraph('Step 1  初步筛选')
doc.add_paragraph('        └── 收集供应商资料（产品目录、认证证书、客户案例）', style='List Bullet 2')
doc.add_paragraph('        └── 完成维度四（资质）的基础核查', style='List Bullet 2')
doc.add_paragraph('        └── 输出：进入深度评估的候选名单', style='List Bullet 2')
doc.add_paragraph()
doc.add_paragraph('Step 2  深度评估')
doc.add_paragraph('        └── 发送结构化问卷（基于上述评估框架）', style='List Bullet 2')
doc.add_paragraph('        └── 安排工厂参观或视频会议深度对接', style='List Bullet 2')
doc.add_paragraph('        └── 索取样品，提交集团技术团队测试', style='List Bullet 2')
doc.add_paragraph('        └── 输出：供应商能力评估报告', style='List Bullet 2')
doc.add_paragraph()
doc.add_paragraph('Step 3  商务谈判')
doc.add_paragraph('        └── 基于评估结果确定合作优先级', style='List Bullet 2')
doc.add_paragraph('        └── 谈判授权、价格、账期、交期等核心条款', style='List Bullet 2')
doc.add_paragraph('        └── 输出：框架合作协议', style='List Bullet 2')
doc.add_paragraph()
doc.add_paragraph('Step 4  试单验证')
doc.add_paragraph('        └── 以小批量订单验证交货能力和产品质量', style='List Bullet 2')
doc.add_paragraph('        └── 建立供应商动态评分档案', style='List Bullet 2')
doc.add_paragraph('        └── 输出：供应商准入确认', style='List Bullet 2')

doc.add_page_break()

add_heading_custom(doc, '第四步：与集团总部技术团队的协同机制', 3)
doc.add_paragraph('鉴于本团队当前技术能力有限，建议建立以下协同机制：')

collab_data = [
    ['协同事项', '频率', '责任方'],
    ['供应商样品技术测试', '按需', '集团技术团队执行，本团队提供样品和测试需求'],
    ['技术参数核实', '按需', '集团技术团队提供判断意见'],
    ['行业标准解读', '月度', '集团技术团队输出简报，本团队学习消化'],
    ['供应商评估结果审核', '每次准入前', '集团技术团队参与最终评审']
]
add_table_with_style(doc, collab_data)

doc.add_page_break()

# 五、线上营销体系建设
add_heading_custom(doc, '五、线上营销体系建设', 1)

add_heading_custom(doc, '5.1 定位原则', 2)
doc.add_paragraph(
    'B2B工业品的线上营销核心不在于流量规模，而在于专业度背书与精准触达。'
    '网站和公众号的首要功能是建立信任，其次才是获客。'
)

add_heading_custom(doc, '5.2 网站专题页', 2)
doc.add_paragraph('上线目标：5月第2周完成').bold = True
doc.add_paragraph('内容框架：')
doc.add_paragraph('• 液冷解决方案定位说明（公司能提供什么）')
doc.add_paragraph('• 产品矩阵展示（快接头、软管系列，含核心技术参数）')
doc.add_paragraph('• 适用场景说明（AIDC、高密度机柜、边缘计算等）')
doc.add_paragraph('• 行业应用案例（初期可使用场景化描述）')
doc.add_paragraph('• 联系与询盘入口')

add_heading_custom(doc, '5.3 公众号内容计划（5月）', 2)
content_data = [
    ['期次', '主题', '目的'],
    ['第1篇', '《AIDC液冷系统对快接头的特殊要求》', '建立技术权威形象'],
    ['第2篇', '《液冷软管选型关键参数解析》', '输出实用干货，吸引技术人员关注'],
    ['第3篇', '《液冷快接头应用案例》', '以案例建立信任'],
    ['第4篇', '《2026液冷市场机会分析》', '展示行业洞察，吸引决策层关注']
]
add_table_with_style(doc, content_data)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('分发策略：').bold = True
p.add_run('公众号发布后同步转发至行业微信群、总经理个人朋友圈，以个人背书强化传播效果。')

doc.add_page_break()

# 六、五月执行里程碑
add_heading_custom(doc, '六、五月执行里程碑', 1)

milestone_data = [
    ['时间节点', '关键任务', '责任人', '完成标志'],
    ['5月第1周', '发布技术工程师招募需求（猎头+直聘）', '总经理/HR', 'JD发布，猎头启动'],
    ['5月第1周', '安排UDQ深度对接会', '总经理', '会议确认，议程发出'],
    ['5月第2周', '完成UDQ对接，形成供应商评估初稿', '总经理', '评估报告完成'],
    ['5月第2周', '网站液冷专题页上线', '内容运营', '页面可访问'],
    ['5月第2周', '公众号第1篇发布', '内容运营', '文章发布'],
    ['5月第3周', '技术工程师候选人面试（不少于3人）', '总经理/HR', '面试记录完成'],
    ['5月第3周', '完成目标客户名单（20家）', '售前项目经理', '名单文档完成'],
    ['5月第3周', '启动前5家目标客户接触', '总经理/售前', '初步沟通记录'],
    ['5月第4周', '技术工程师到位或顾问协议签署', '总经理/HR', '合同/协议签署'],
    ['5月第4周', '完成售前SOP文档', '售前项目经理', '文档审核通过'],
    ['5月第4周', '月度复盘，制定6月计划', '总经理', '复盘会议完成']
]
add_table_with_style(doc, milestone_data)

doc.add_page_break()

# 七、风险提示与应对
add_heading_custom(doc, '七、风险提示与应对', 1)

risk_data = [
    ['风险项', '风险描述', '应对措施'],
    ['技术人员招募周期长', '液冷技术人才稀缺，5月内到岗存在不确定性', '同步推进顾问/兼职合作，借助集团技术团队过渡'],
    ['供应商评估能力不足', '团队技术背景薄弱，难以独立判断产品技术真实性', '强化与集团技术团队协同，引入第三方检测机构'],
    ['三线并行资源分散', '同时推进团队、营销、供应链，精力有限', '优先级排序：供应链稳定 > 技术人员到位 > 线上营销'],
    ['客户转化周期长', 'B2B工业品决策链长，5月内难以形成订单', '5月目标定位为"建立商机管道"，不以签单为考核']
]
add_table_with_style(doc, risk_data)

doc.add_page_break()

# 八、附件说明
add_heading_custom(doc, '八、附件说明', 1)
doc.add_paragraph('• 附件一：技术工程师岗位JD（可直接提供给HR/猎头）')
doc.add_paragraph('• 附件二：供应商评估问卷模板')
doc.add_paragraph('• 附件三：目标客户名单模板')
doc.add_paragraph()
doc.add_paragraph('以上附件可根据需要单独输出。').italic = True

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('文件结束').bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
p = doc.add_paragraph('本文件为内部工作文件，请妥善保管。')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].italic = True
p.runs[0].font.size = Pt(9)

# 保存文档
output_path = '/Users/landylook/简报网站/docs/液冷市场突破计划-2026年5月.docx'
doc.save(output_path)
print(f"Word文档已生成：{output_path}")
